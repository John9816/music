# -*- coding: utf-8 -*-
import os
import re
import csv
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def parse_num(s):
    if s is None:
        return None
    t = (
        s.replace(",", "")
        .replace("，", "")
        .replace(" ", "")
        .replace("\u3000", "")
        .replace("\u00a0", "")
        .strip()
    )
    if not t:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", t)
    return float(m.group()) if m else None


def fmt(v):
    return f"{v:.2f}"


def clean(s):
    return " ".join(s.replace("\n", " ").split())


src = os.environ["DOCX_PATH"]
doc = Document(src)
body = doc.element.body

# 把每个表格归属到它所在的专业段落（工程名称）
table_sections = {}
section = "未知"
ti = 0
for child in body.iterchildren():
    if child.tag == qn("w:tbl"):
        table_sections[ti] = section
        ti += 1
    elif child.tag == qn("w:p"):
        p = Paragraph(child, doc)
        t = p.text.strip()
        if t.startswith("工程名称"):
            m = re.search(r"工程名称\s*[:：]\s*(\S+)", t)
            if m:
                section = m.group(1).strip()

rows_out = []
sec_running = {}
sec_count = {}

for tidx in range(len(doc.tables)):
    table = doc.tables[tidx]
    # 定位表头行（含“综合单价”“合价”的 8 列表）
    hdr = None
    is_measure = False
    for ri, row in enumerate(table.rows[:3]):
        texts = [c.text for c in row.cells]
        if len(row.cells) >= 8 and any("综合单价" in t for t in texts) and any("合" in t for t in texts):
            hdr = ri
            if any("项目特征或工作内容" in t for t in texts):
                is_measure = True
            break
    if hdr is None:
        continue

    sec = table_sections.get(tidx, "未知")
    table_total = 0.0
    table_count = 0

    for ri in range(hdr + 1, len(table.rows)):
        cells = table.rows[ri].cells
        if len(cells) < 8:
            continue
        texts = [c.text for c in cells]
        seq = texts[0].strip().replace("\n", "")
        if seq == "序号" or (seq and not re.search(r"\d", seq)):
            # 章节行（如“拆除工程”“A.5 混凝土…”）跳过；带序号的数字行继续
            code_tmp = texts[1].strip()
            if not re.search(r"\d", code_tmp):
                continue
        price = parse_num(texts[6])
        if price is None:
            continue
        qty = parse_num(texts[5])
        if qty is None:
            continue
        total_i = round(qty * price, 2)
        table_total += total_i
        table_count += 1
        cells[7].text = fmt(total_i)
        rows_out.append(
            [
                tidx,
                ri,
                sec,
                seq,
                texts[1].strip().replace("\n", ""),
                clean(texts[2]),
                texts[4].strip(),
                qty,
                price,
                total_i,
            ]
        )

    sec_running[sec] = sec_running.get(sec, 0.0) + table_total
    sec_count[sec] = sec_count.get(sec, 0) + table_count

    # 分部分项表底部的“合 计”行：填该专业累计合价（措施表不填）
    if not is_measure and table_count > 0:
        for ri in range(len(table.rows) - 1, hdr, -1):
            cells = table.rows[ri].cells
            if (
                len(cells) >= 8
                and "合" in cells[2].text
                and "计" in cells[3].text
            ):
                cells[7].text = fmt(sec_running[sec])
                break

out_dir = os.path.dirname(src)
out_name = os.path.splitext(os.path.basename(src))[0] + "_已合计.docx"
out_path = os.path.join(out_dir, out_name)
doc.save(out_path)

csv_path = r"D:\Projects\music\_hejia_details.csv"
with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
    w = csv.writer(f)
    w.writerow(
        ["表号", "行号", "专业", "序号", "项目编码", "项目名称", "单位", "工程数量", "综合单价", "合价"]
    )
    w.writerows(rows_out)

grand = sum(r[9] for r in rows_out)
print("已写入文件:", out_path)
print("已计算合价的行数:", len(rows_out))
for sec in sec_running:
    print(f"{sec}: 行数={sec_count[sec]} 合计={fmt(sec_running[sec])}")
print("总计:", fmt(grand))
print("明细CSV:", csv_path)
